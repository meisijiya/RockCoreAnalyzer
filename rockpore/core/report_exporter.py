"""统一报告导出器.

v1.2.0 新增: 三个模块(孔洞/裂缝/粒度)共享报告导出逻辑.
支持格式:
- HTML  (内置, 无依赖)
- TXT   (内置, 无依赖)
- PDF   (reportlab)
- XLSX  (openpyxl)
- DOCX  (python-docx)
- CSV   (内置, 无依赖, 只导出数据表)

Usage:
    exporter = ReportExporter(
        title="孔洞分析报告",
        info={"井号": "A1", "深度": "1234.5 m", ...},
        summary={"孔洞总数": 12, "面孔率": 5.2, ...},
        headers=["ID", "面积(mm²)", "直径(mm)", ...],
        rows=[[1, 1.5, 1.4, ...], [2, 2.3, 1.7, ...], ...],
    )
    exporter.export("html", "report.html")
    exporter.export("pdf", "report.pdf")
"""
from __future__ import annotations

import csv
import os
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Sequence, Union

import numpy as np


# 格式元信息 (扩展名, MIME)
SUPPORTED_FORMATS: Dict[str, Dict[str, str]] = {
    "html": {"ext": "html", "label": "HTML 网页 (*.html)", "mime": "text/html"},
    "txt":  {"ext": "txt",  "label": "纯文本 (*.txt)",      "mime": "text/plain"},
    "pdf":  {"ext": "pdf",  "label": "PDF 文档 (*.pdf)",     "mime": "application/pdf"},
    "xlsx": {"ext": "xlsx", "label": "Excel 表格 (*.xlsx)",  "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"},
    "docx": {"ext": "docx", "label": "Word 文档 (*.docx)",   "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    "csv":  {"ext": "csv",  "label": "CSV 表格 (*.csv)",     "mime": "text/csv"},
}


@dataclass
class ReportExporter:
    """统一报告导出器.

    Attributes:
        title: 报告标题 (如 "孔洞分析报告")
        info: 项目元数据 (key: value)
        summary: 摘要统计 (key: value)
        headers: 表格表头
        rows: 表格数据行
        annotated_image: 标注图像 (BGR ndarray, 可选)
        notes: 备注文字 (可选, 多行)
    """
    title: str
    info: Dict[str, Any] = field(default_factory=dict)
    summary: Dict[str, Any] = field(default_factory=dict)
    headers: List[str] = field(default_factory=list)
    rows: List[List[Any]] = field(default_factory=list)
    annotated_image: Optional[np.ndarray] = None
    notes: str = ""

    def _format_value(self, v: Any) -> str:
        """格式化单个值 (处理 numpy 类型等)."""
        if isinstance(v, float):
            if abs(v) < 0.01 or abs(v) > 1e6:
                return f"{v:.3e}"
            return f"{v:.3f}".rstrip("0").rstrip(".") or "0"
        if isinstance(v, (np.integer,)):
            return str(int(v))
        if isinstance(v, (np.floating,)):
            return self._format_value(float(v))
        if isinstance(v, np.ndarray):
            return str(v.tolist())
        return str(v)

    def _build_text(self) -> str:
        """生成纯文本报告."""
        lines = []
        lines.append("═" * 60)
        lines.append(f"  {self.title}")
        lines.append("═" * 60)
        lines.append("")
        if self.info:
            lines.append("── 项目信息 ──")
            for k, v in self.info.items():
                lines.append(f"  {k}: {v}")
            lines.append("")
        if self.summary:
            lines.append("── 摘要统计 ──")
            for k, v in self.summary.items():
                lines.append(f"  {k}: {v}")
            lines.append("")
        if self.headers and self.rows:
            lines.append("── 详细数据 ──")
            lines.append("  " + " | ".join(self.headers))
            lines.append("  " + "-+-".join("-" * len(h) for h in self.headers))
            for row in self.rows:
                lines.append("  " + " | ".join(self._format_value(v) for v in row))
            lines.append("")
        if self.notes:
            lines.append("── 备注 ──")
            lines.append(self.notes)
        return "\n".join(lines)

    def _build_html(self) -> str:
        """生成 HTML 报告 (BGR 图像嵌入 base64)."""
        import base64
        img_html = ""
        if self.annotated_image is not None:
            import cv2
            _, buf = cv2.imencode(".png", self.annotated_image)
            b64 = base64.b64encode(buf.tobytes()).decode("ascii")
            img_html = (
                f'<div style="margin: 20px 0;">'
                f'<img src="data:image/png;base64,{b64}" '
                f'style="max-width: 100%; border: 1px solid #ddd; border-radius: 4px;"/>'
                f'</div>'
            )
        info_rows = "".join(
            f"<tr><th>{self._format_value(k)}</th><td>{self._format_value(v)}</td></tr>"
            for k, v in self.info.items()
        )
        summary_rows = "".join(
            f"<tr><th>{self._format_value(k)}</th><td>{self._format_value(v)}</td></tr>"
            for k, v in self.summary.items()
        )
        if self.headers and self.rows:
            head = "".join(f"<th>{h}</th>" for h in self.headers)
            body = "".join(
                "<tr>" + "".join(f"<td>{self._format_value(v)}</td>" for v in row) + "</tr>"
                for row in self.rows
            )
            table = f'<table border="1" cellpadding="6" cellspacing="0" style="border-collapse: collapse; width: 100%;"><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>'
        else:
            table = ""
        notes_html = ""
        if self.notes:
            notes_html = f'<h2>备注</h2><pre style="background: #f5f5f5; padding: 12px; border-radius: 4px;">{self.notes}</pre>'
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{self.title}</title>
<style>
body {{ font-family: -apple-system, "Segoe UI", "Microsoft YaHei", sans-serif; margin: 30px; color: #1f2328; }}
h1 {{ color: #2c5fa3; border-bottom: 2px solid #2c5fa3; padding-bottom: 8px; }}
h2 {{ color: #2c5fa3; margin-top: 24px; }}
table {{ margin: 12px 0; font-size: 13px; }}
th {{ background: #e8f0fb; text-align: left; padding: 6px 12px; }}
td {{ padding: 6px 12px; }}
.info-table th {{ width: 120px; }}
.summary-table th {{ width: 160px; }}
.notes {{ background: #f5f5f5; padding: 12px; border-radius: 4px; }}
</style>
</head>
<body>
<h1>{self.title}</h1>
{img_html}
{f'<h2>项目信息</h2><table class="info-table">{info_rows}</table>' if self.info else ''}
{f'<h2>摘要统计</h2><table class="summary-table">{summary_rows}</table>' if self.summary else ''}
{f'<h2>详细数据</h2>{table}' if table else ''}
{notes_html}
<hr style="margin-top: 32px; color: #ddd;"/>
<p style="color: #8b949e; font-size: 11px;">岩心图像分析软件 v1.2.0  ·  生成时间: {self._now()}</p>
</body>
</html>"""

    def _now(self) -> str:
        """当前时间字符串."""
        import datetime
        return datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def export(self, fmt: str, output_path: str) -> bool:
        """导出报告.

        Args:
            fmt: 格式 ("html"/"txt"/"pdf"/"xlsx"/"docx"/"csv")
            output_path: 输出文件路径 (不含扩展名也行,会自动加)

        Returns:
            True 成功 / False 失败
        """
        fmt = fmt.lower().strip()
        if fmt not in SUPPORTED_FORMATS:
            raise ValueError(f"不支持的格式: {fmt}, 必须是 {list(SUPPORTED_FORMATS)}")
        # 自动补扩展名
        expected_ext = "." + SUPPORTED_FORMATS[fmt]["ext"]
        if not output_path.lower().endswith(expected_ext):
            output_path = output_path + expected_ext
        os.makedirs(os.path.dirname(os.path.abspath(output_path)) or ".", exist_ok=True)
        try:
            if fmt == "html":
                content = self._build_html()
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(content)
            elif fmt == "txt":
                content = self._build_text()
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(content)
            elif fmt == "csv":
                with open(output_path, "w", encoding="utf-8", newline="") as f:
                    w = csv.writer(f)
                    w.writerow([self.title])
                    for k, v in self.info.items():
                        w.writerow([k, v])
                    if self.summary:
                        w.writerow([])
                        w.writerow(["摘要统计"])
                        for k, v in self.summary.items():
                            w.writerow([k, v])
                    if self.headers:
                        w.writerow([])
                        w.writerow(self.headers)
                        w.writerows(self.rows)
            elif fmt == "pdf":
                self._export_pdf(output_path)
            elif fmt == "xlsx":
                self._export_xlsx(output_path)
            elif fmt == "docx":
                self._export_docx(output_path)
            return True
        except Exception as e:
            # 重新抛,让调用方知道
            raise RuntimeError(f"导出 {fmt} 失败: {e}") from e

    def _export_pdf(self, output_path: str):
        """PDF 导出 (使用 reportlab)."""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image,
        )
        # 注册中文字体 (尝试常见路径)
        font_name = self._register_chinese_font()
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("TitleCN", parent=styles["Title"], fontName=font_name, fontSize=18)
        h2_style = ParagraphStyle("H2CN", parent=styles["Heading2"], fontName=font_name, fontSize=13)
        body_style = ParagraphStyle("BodyCN", parent=styles["BodyText"], fontName=font_name, fontSize=10)
        doc = SimpleDocTemplate(output_path, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        story = []
        story.append(Paragraph(self.title, title_style))
        story.append(Spacer(1, 0.5*cm))
        if self.annotated_image is not None:
            import tempfile, cv2
            tmp_png = os.path.join(tempfile.gettempdir(), "_report_img.png")
            cv2.imwrite(tmp_png, self.annotated_image)
            img = Image(tmp_png, width=15*cm, height=10*cm, kind="proportional")
            story.append(img)
            story.append(Spacer(1, 0.5*cm))
        if self.info:
            story.append(Paragraph("项目信息", h2_style))
            data = [[Paragraph(self._format_value(k), body_style), Paragraph(self._format_value(v), body_style)] for k, v in self.info.items()]
            t = Table(data, colWidths=[4*cm, 12*cm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e8f0fb")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#ccc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(t)
            story.append(Spacer(1, 0.3*cm))
        if self.summary:
            story.append(Paragraph("摘要统计", h2_style))
            data = [[Paragraph(self._format_value(k), body_style), Paragraph(self._format_value(v), body_style)] for k, v in self.summary.items()]
            t = Table(data, colWidths=[4*cm, 12*cm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e8f0fb")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#ccc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(t)
            story.append(Spacer(1, 0.3*cm))
        if self.headers and self.rows:
            story.append(Paragraph("详细数据", h2_style))
            data = [[Paragraph(self._format_value(v), body_style) for v in row] for row in self.rows]
            data.insert(0, [Paragraph(h, body_style) for h in self.headers])
            t = Table(data)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c5fa3")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#ccc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(t)
        if self.notes:
            story.append(Spacer(1, 0.3*cm))
            story.append(Paragraph("备注", h2_style))
            for line in self.notes.splitlines():
                story.append(Paragraph(line, body_style))
        doc.build(story)

    def _export_xlsx(self, output_path: str):
        """Excel 导出 (使用 openpyxl)."""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.drawing.image import Image as XLImage
        wb = Workbook()
        # Sheet 1: 报告
        ws = wb.active
        ws.title = "报告"
        # 标题
        ws["A1"] = self.title
        ws["A1"].font = Font(bold=True, size=16, color="2C5FA3")
        row_idx = 3
        if self.info:
            ws[f"A{row_idx}"] = "项目信息"
            ws[f"A{row_idx}"].font = Font(bold=True, size=12, color="2C5FA3")
            row_idx += 1
            for k, v in self.info.items():
                ws[f"A{row_idx}"] = k
                ws[f"B{row_idx}"] = self._format_value(v)
                row_idx += 1
            row_idx += 1
        if self.summary:
            ws[f"A{row_idx}"] = "摘要统计"
            ws[f"A{row_idx}"].font = Font(bold=True, size=12, color="2C5FA3")
            row_idx += 1
            for k, v in self.summary.items():
                ws[f"A{row_idx}"] = k
                ws[f"B{row_idx}"] = self._format_value(v)
                row_idx += 1
            row_idx += 1
        # 详细数据 sheet
        if self.headers and self.rows:
            ws2 = wb.create_sheet("详细数据")
            for col_idx, h in enumerate(self.headers, start=1):
                cell = ws2.cell(row=1, column=col_idx, value=h)
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="2C5FA3")
            for r_idx, row in enumerate(self.rows, start=2):
                for c_idx, v in enumerate(row, start=1):
                    ws2.cell(row=r_idx, column=c_idx, value=self._format_value(v))
        # 标注图
        if self.annotated_image is not None:
            import tempfile
            tmp_png = os.path.join(tempfile.gettempdir(), "_report_img.png")
            import cv2
            cv2.imwrite(tmp_png, self.annotated_image)
            img = XLImage(tmp_png)
            img.width = 480
            img.height = 320
            ws.add_image(img, f"A{row_idx + 2}")
        wb.save(output_path)

    def _export_docx(self, output_path: str):
        """Word 导出 (使用 python-docx)."""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        # 标题
        title = doc.add_heading(self.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 标注图
        if self.annotated_image is not None:
            import tempfile, cv2
            tmp_png = os.path.join(tempfile.gettempdir(), "_report_img.png")
            cv2.imwrite(tmp_png, self.annotated_image)
            doc.add_picture(tmp_png, width=Inches(6))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if self.info:
            doc.add_heading("项目信息", level=1)
            t = doc.add_table(rows=len(self.info), cols=2)
            t.style = "Light Grid Accent 1"
            for i, (k, v) in enumerate(self.info.items()):
                t.cell(i, 0).text = self._format_value(k)
                t.cell(i, 1).text = self._format_value(v)
        if self.summary:
            doc.add_heading("摘要统计", level=1)
            t = doc.add_table(rows=len(self.summary), cols=2)
            t.style = "Light Grid Accent 1"
            for i, (k, v) in enumerate(self.summary.items()):
                t.cell(i, 0).text = self._format_value(k)
                t.cell(i, 1).text = self._format_value(v)
        if self.headers and self.rows:
            doc.add_heading("详细数据", level=1)
            t = doc.add_table(rows=len(self.rows) + 1, cols=len(self.headers))
            t.style = "Light Grid Accent 1"
            for j, h in enumerate(self.headers):
                cell = t.cell(0, j)
                cell.text = h
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for i, row in enumerate(self.rows, start=1):
                for j, v in enumerate(row):
                    t.cell(i, j).text = self._format_value(v)
        if self.notes:
            doc.add_heading("备注", level=1)
            doc.add_paragraph(self.notes)
        doc.save(output_path)

    def _register_chinese_font(self) -> str:
        """注册中文字体供 reportlab 使用,返回字体名."""
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        # 常见中文字体路径
        candidates = [
            # Linux
            ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", "WenQuanYiMicroHei"),
            ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", "WenQuanYiZenHei"),
            ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK"),
            ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK"),
            # Windows
            ("C:/Windows/Fonts/msyh.ttc", "MicrosoftYaHei"),
            ("C:/Windows/Fonts/simhei.ttf", "SimHei"),
            ("C:/Windows/Fonts/simsun.ttc", "SimSun"),
        ]
        for path, name in candidates:
            if os.path.exists(path):
                try:
                    pdfmetrics.registerFont(TTFont(name, path))
                    return name
                except Exception:
                    continue
        # fallback: 用 reportlab 自带字体 (不支持中文,但不报错)
        return "Helvetica"


__all__ = ["ReportExporter", "SUPPORTED_FORMATS"]